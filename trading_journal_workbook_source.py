# -*- coding: utf-8 -*-
"""
Trading Notes — Daily Trading Journal (30 Days of Process)
Source code reconstructed from:
    Trading_Notes_Daily_Journal_30_Days.pdf

Requirements:
    pip install python-docx

Output:
    Trading_Notes_Daily_Journal_30_Days_generated.docx

Notes:
- A4 portrait
- Clean / print-friendly theme
- 57 pages
- Header: TRADING NOTES | Process Before Profit
- 30 Daily Journal pages
- System-neutral workbook
"""

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = "Trading_Notes_Daily_Journal_30_Days_generated.docx"

# ---------------------------------------------------------------------
# THEME
# ---------------------------------------------------------------------

A4_W = 210
A4_H = 297

MARGIN_TOP = 15
MARGIN_BOTTOM = 15
MARGIN_LEFT = 16
MARGIN_RIGHT = 16

FONT = "Arial"
BODY = 10.8
SMALL = 8.2
LABEL = 9.2
H1 = 19
H2 = 12.5
TITLE = 28

THIN_GRAY = "B8B8B8"

doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(A4_W)
sec.page_height = Mm(A4_H)
sec.top_margin = Mm(MARGIN_TOP)
sec.bottom_margin = Mm(MARGIN_BOTTOM)
sec.left_margin = Mm(MARGIN_LEFT)
sec.right_margin = Mm(MARGIN_RIGHT)

styles = doc.styles
styles["Normal"].font.name = FONT
styles["Normal"].font.size = Pt(BODY)

for name, size, bold in [
    ("TitleX", TITLE, True),
    ("H1X", H1, True),
    ("H2X", H2, True),
    ("SmallX", SMALL, False),
]:
    if name not in styles:
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles[name]
    st.font.name = FONT
    st.font.size = Pt(size)
    st.font.bold = bold


# ---------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------

def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=THIN_GRAY, size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        e = borders.find(qn(tag))
        if e is None:
            e = OxmlElement(tag)
            borders.append(e)
        e.set(qn("w:val"), "nil")


def add_bottom_rule(paragraph, color="A0A0A0", size="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:color"), color)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)


def p(text="", style=None, bold=False, size=None, align=None,
      before=0, after=5, line=1.08):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = FONT
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    return para


def page_break():
    doc.add_page_break()


def page_header(page_no):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Mm(95)
    table.columns[1].width = Mm(85)
    remove_table_borders(table)

    left, right = table.rows[0].cells
    left.text = "TRADING NOTES  |  Process Before Profit"
    right.text = f"{page_no:02d}"
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for cell in (left, right):
        set_cell_margins(cell, 0, 0, 0, 0)
        for run in cell.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(7.2)
            run.bold = run.text.startswith("TRADING NOTES")

    rule = p("", after=7)
    add_bottom_rule(rule, "B0B0B0", "3")


def footer():
    for section in doc.sections:
        ft = section.footer
        para = ft.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("TRADING NOTES  •  DAILY TRADING JOURNAL  •  PROCESS BEFORE PROFIT")
        run.font.name = FONT
        run.font.size = Pt(6.8)


def title(en, th=None):
    p(en, "H1X", after=1)
    if th:
        p(th, "H2X", after=9)


def writing_lines(n=2, width=76):
    for _ in range(n):
        p("_" * width, size=9.2, after=6)


def scale_row(table, row_index, height_mm):
    row = table.rows[row_index]
    row.height = Mm(height_mm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


# ---------------------------------------------------------------------
# PAGE 01 — COVER
# ---------------------------------------------------------------------

page_header(1)
p("TRADING NOTES", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, before=22, after=25)
p("DAILY", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p("TRADING JOURNAL", bold=True, size=27, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
p("30 Days of Process", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=35)
p("PROCESS BEFORE PROFIT", bold=True, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
p("กระบวนการมาก่อนผลลัพธ์", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# PAGE 02 — BELONGS TO
# ---------------------------------------------------------------------

page_header(2)
title("THIS WORKBOOK BELONGS TO", "Workbook เล่มนี้เป็นของ")
for label in ["NAME", "START DATE", "END DATE", "MY 30-DAY FOCUS"]:
    p(label, bold=True, size=LABEL, before=6, after=1)
    writing_lines(1)
page_break()


# ---------------------------------------------------------------------
# PAGE 03 — DISCLAIMER
# ---------------------------------------------------------------------

page_header(3)
title("DISCLAIMER", "ข้อจำกัดความรับผิดชอบ")
disc = [
    "เนื้อหาและแบบฝึกหัดทั้งหมดใน Workbook เล่มนี้จัดทำขึ้นเพื่อวัตถุประสงค์ด้านการศึกษา "
    "การเรียนรู้ และการพัฒนากระบวนการตัดสินใจในการเทรดเท่านั้น",
    "เนื้อหาทั้งหมดไม่ถือเป็นคำแนะนำด้านการลงทุน การเงิน หรือการซื้อขายหลักทรัพย์ "
    "สินทรัพย์ หรือผลิตภัณฑ์ทางการเงินใด ๆ",
    "Workbook นี้ไม่มีการให้สัญญาณซื้อขาย ไม่มีการรับประกันผลตอบแทน "
    "และไม่ได้มีวัตถุประสงค์เพื่อชี้นำให้ซื้อหรือขายสินทรัพย์ใด",
    "การลงทุนและการเทรดมีความเสี่ยง ผู้ใช้งานควรศึกษาข้อมูล ประเมินความเสี่ยง "
    "และรับผิดชอบต่อการตัดสินใจของตนเอง",
]
for text in disc:
    p(text, size=10.7, after=11, line=1.18)
page_break()


# ---------------------------------------------------------------------
# PAGE 04 — CONTENTS
# ---------------------------------------------------------------------

page_header(4)
title("CONTENTS", "สารบัญ")

p("FOUNDATION", bold=True, size=11, after=4)
contents = [
    ("01", "Welcome", "ก่อนเริ่มต้น Journal เล่มนี้"),
    ("02", "Why Trading Journal?", "ทำไม Trader ควรมี Journal"),
    ("03", "Process Before Profit", "กระบวนการมาก่อนผลลัพธ์"),
    ("04", "Good Trade ≠ Winning Trade", "Trade ที่ดีไม่ได้หมายถึง Trade ที่กำไร"),
    ("05", "Know Yourself", "ก่อนวิเคราะห์ตลาด ลองวิเคราะห์ตัวเอง"),
    ("06", "Define Your Trading System", "นิยามระบบการเทรดของตัวเอง"),
    ("07", "Define Your A+ Setup", "นิยาม Setup ที่ดีที่สุดของตัวเอง"),
    ("08", "My No-Trade Rules", "เมื่อไรที่ฉันจะไม่เทรด"),
    ("09", "How to Use This Journal", "วิธีใช้ Journal ตลอด 30 วัน"),
]
for no, en, th in contents:
    p(f"{no}   {en}   {th}", size=9.5, after=3)

p("30 DAYS OF PROCESS", bold=True, size=11, before=7, after=2)
p("Daily Trading Journal   Day 01–30", size=9.5, after=6)

p("END OF 30 DAYS", bold=True, size=11, before=5, after=2)
p("Looking Back • Process Check • Day 1 vs Day 30 • Next Step", size=9.5)
page_break()


# ---------------------------------------------------------------------
# PAGE 05 — WELCOME
# ---------------------------------------------------------------------

page_header(5)
title("01 — WELCOME", "ก่อนเริ่มต้น Journal เล่มนี้")

for text in [
    "การพัฒนาการเทรดไม่ได้เกิดจากการหาจุดเข้าให้แม่นขึ้นเพียงอย่างเดียว",
    "หลายครั้ง Trader รู้วิธีวิเคราะห์กราฟ รู้จัก Setup รู้ว่าควรวาง Stop Loss ที่ไหน "
    "และรู้ว่าควรบริหารความเสี่ยงอย่างไร แต่เมื่ออยู่หน้ากราฟจริง "
    "เราอาจตัดสินใจแตกต่างจากสิ่งที่ตัวเองรู้",
    "เราอาจเข้าเร็วเพราะกลัวตกรถ เข้า Trade ที่ไม่ได้อยู่ในแผน เพิ่มความเสี่ยงหลังขาดทุน "
    "ปิดกำไรเร็ว ขยับ Stop Loss หรือกด Order เพียงเพราะอยากมีส่วนร่วมกับตลาด",
]:
    p(text, size=10.8, after=10, line=1.2)

p("“สิ่งที่เรารู้ กับสิ่งที่เราทำจริง เหมือนกันหรือไม่?”",
  bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=14)

p("Trading Journal เล่มนี้ถูกสร้างขึ้นเพื่อช่วยให้เราเห็นช่องว่างตรงนั้น",
  size=10.8, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# PAGE 06 — WHY JOURNAL
# ---------------------------------------------------------------------

page_header(6)
title("02 — WHY TRADING JOURNAL?", "ทำไม Trader ควรมี Journal")

paras = [
    "Trade History สามารถบอกเราได้ว่าเรา Buy หรือ Sell ที่ราคาไหน ได้กำไรหรือขาดทุนเท่าไร "
    "แต่ตัวเลขเหล่านั้นไม่ได้บอกว่า “ทำไมเราถึงเข้า Trade นั้น”",
    "มันไม่ได้บอกว่าเราทำตามแผนหรือไม่ กำลัง FOMO หรือไม่ เพิ่ม Risk หลังขาดทุนหรือไม่ "
    "และไม่ได้บอกว่า Trade ที่ขาดทุนอาจเป็น Trade ที่ทำถูกต้องทุกอย่างแล้ว",
]
for x in paras:
    p(x, size=10.8, after=10, line=1.18)

p("Trade History บันทึก “ผล”", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=4)
p("แต่ Trading Journal บันทึก “กระบวนการ”", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=15)

p("เมื่อบันทึกอย่างต่อเนื่อง เราจะเริ่มเห็นว่าอะไรที่เราทำได้ดีซ้ำ ๆ "
  "และอะไรที่เราทำผิดซ้ำ ๆ", size=10.8, after=10)
p("Journal ไม่ได้มีไว้เพื่อตำหนิตัวเอง แต่มีไว้เพื่อสร้างข้อมูลที่ช่วยให้เราเข้าใจ "
  "Trader คนหนึ่งให้มากขึ้น", size=10.8, after=8)
p("Trader คนนั้นก็คือตัวเราเอง", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# PAGE 07 — PROCESS BEFORE PROFIT
# ---------------------------------------------------------------------

page_header(7)
title("03 — PROCESS BEFORE PROFIT", "กระบวนการมาก่อนผลลัพธ์")

p("หนึ่งในกับดักที่พบได้บ่อยคือการคิดว่า", size=10.8, after=4)
p("กำไร = ทำถูก", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p("ขาดทุน = ทำผิด", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
p("แต่ในความเป็นจริงไม่ได้เป็นเช่นนั้นเสมอไป", size=10.8, after=12)

p("TRADE A", bold=True, size=11)
p("ไม่มี Setup ชัดเจน • เข้าเพราะ FOMO • Risk มากกว่าที่กำหนด • สุดท้ายกำไร +3R",
  size=10.3, after=9)

p("TRADE B", bold=True, size=11)
p("Setup ตรงตาม Trading Plan • รอ Confirmation • Risk ตามกฎ • วาง Stop Loss ตามแผน • สุดท้ายขาดทุน -1R",
  size=10.3, after=11)

p("ถ้าดูเฉพาะ P/L Trade A ดูเหมือนดีกว่า แต่ถ้าถามว่า "
  "“Trade แบบไหนที่เราต้องการทำซ้ำอีก 100 ครั้ง?” คำตอบอาจเปลี่ยนไป",
  size=10.8, after=11)

p("Good Process ไม่ได้รับประกันว่า Trade นั้นจะกำไร "
  "แต่เป็นสิ่งที่เราสามารถทำซ้ำ วัดผล และพัฒนาต่อได้",
  size=10.8, after=15)

p("Don't judge the process by one outcome.",
  bold=True, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# PAGE 08 — GOOD TRADE ≠ WINNING TRADE
# ---------------------------------------------------------------------

page_header(8)
title("04 — GOOD TRADE ≠ WINNING TRADE", "Trade ที่ดี ไม่ได้หมายถึง Trade ที่กำไร")

matrix = doc.add_table(rows=3, cols=3)
matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
matrix.style = "Table Grid"
vals = [
    ["PROCESS", "WIN", "LOSS"],
    ["FOLLOWED THE PROCESS", "GOOD TRADE", "GOOD TRADE"],
    ["BROKE THE PROCESS", "BAD TRADE", "BAD TRADE"],
]
for i, row in enumerate(vals):
    for j, val in enumerate(row):
        c = matrix.cell(i, j)
        c.text = val
        set_cell_margins(c, 90, 80, 90, 80)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in c.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(9)
            run.bold = True
        set_cell_border(c)
    scale_row(matrix, i, 11)

items = [
    ("GOOD TRADE + WIN", "ทำตามแผนและกำไร", "เราทำสิ่งที่ควรทำ และตลาดให้ผลตอบแทน"),
    ("GOOD TRADE + LOSS", "ทำตามแผนแต่ขาดทุน", "Setup, Risk และ Execution ถูกต้อง แต่ตลาดไม่ได้เคลื่อนไหวตาม Scenario"),
    ("BAD TRADE + WIN", "ผิดแผนแต่กำไร", "กำไรอาจทำให้พฤติกรรมที่ผิดได้รับการตอกย้ำ จึงต้องระวังเป็นพิเศษ"),
    ("BAD TRADE + LOSS", "ผิดแผนและขาดทุน", "สิ่งที่ต้อง Review คือ “ทำไมเราจึงตัดสินใจออกนอก Process?”"),
]
for en, th, desc in items:
    p(f"{en} — {th}", bold=True, size=10.3, before=7, after=2)
    p(desc, size=9.8, after=2)
page_break()


# ---------------------------------------------------------------------
# PAGE 09 — KNOW YOURSELF / SELF-ASSESSMENT
# ---------------------------------------------------------------------

ASSESSMENT_ITEMS = [
    "ฉันมี Trading System ที่อธิบายได้ชัดเจน",
    "ฉันรู้ว่า Setup แบบไหนคือ Setup หลักของตัวเอง",
    "ฉันรู้ว่าเมื่อไรควร Trade และเมื่อไรควร No Trade",
    "ฉันกำหนด Risk ก่อนเข้า Order",
    "ฉันสามารถรอ Confirmation ตามแผนได้",
    "ฉันยอมรับ Stop Loss โดยไม่พยายามแก้มือทันที",
    "ฉันสามารถหยุดเทรดเมื่อไม่มี Setup",
    "ฉันแยกได้ว่ากำลังเข้าเพราะ Setup หรือ Emotion",
    "ฉัน Review Trade ของตัวเองอย่างสม่ำเสมอ",
]

page_header(9)
title("05 — KNOW YOURSELF", "ก่อนวิเคราะห์ตลาด ลองวิเคราะห์ตัวเอง")
p("ก่อนเริ่ม Day 1 ให้ประเมินตัวเองจากพฤติกรรมจริง ไม่มีคำตอบถูกหรือผิด",
  size=10.5, after=5)
p("1 = ยังไม่ชัดเจน / ทำไม่ได้สม่ำเสมอ      5 = ชัดเจน / ทำได้สม่ำเสมอ",
  bold=True, size=9.5, after=8)

assess = doc.add_table(rows=len(ASSESSMENT_ITEMS), cols=2)
assess.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(assess)
for i, item in enumerate(ASSESSMENT_ITEMS):
    assess.cell(i,0).text = item
    assess.cell(i,1).text = "1    2    3    4    5"
    assess.cell(i,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for j in (0,1):
        set_cell_margins(assess.cell(i,j), 65, 20, 65, 20)
        for run in assess.cell(i,j).paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(9.2)
    scale_row(assess, i, 8.5)
page_break()


# ---------------------------------------------------------------------
# PAGE 10 — SELF-ASSESSMENT REFLECTION
# ---------------------------------------------------------------------

page_header(10)
title("SELF-ASSESSMENT REFLECTION", "ทบทวนตัวเองก่อนเริ่ม")

for en, th in [
    ("MY CURRENT STRENGTH", "สิ่งที่ฉันทำได้ดีที่สุดในตอนนี้"),
    ("MY MAIN AREA TO IMPROVE", "สิ่งที่ฉันต้องการพัฒนามากที่สุด"),
    ("AFTER 30 DAYS, I WANT TO...", "ในอีก 30 วัน ฉันอยากเห็นตัวเองเปลี่ยนแปลงเรื่องใด"),
]:
    p(f"{en}   {th}", bold=True, size=10.2, before=8, after=2)
    writing_lines(2)

p("อย่าพยายามเป็น Trader คนใหม่ในวันเดียว\nเพียงเริ่มจากการมองเห็น Trader ที่เราเป็นอยู่ในวันนี้",
  bold=True, size=11.2, align=WD_ALIGN_PARAGRAPH.CENTER, before=12)
page_break()


# ---------------------------------------------------------------------
# PAGE 11 — DEFINE YOUR TRADING SYSTEM
# ---------------------------------------------------------------------

page_header(11)
title("06 — DEFINE YOUR TRADING SYSTEM", "นิยามระบบการเทรดของตัวเอง")
p("Journal จะวัด Process ไม่ได้ หากเรายังไม่มีคำจำกัดความว่า Process ของเราคืออะไร",
  size=10.8, after=10)
p("Trading System ไม่จำเป็นต้องซับซ้อน แต่ต้องชัดเจนพอที่จะตอบได้ว่า "
  "“อะไรคือ Trade ที่อยู่ในระบบของฉัน?”",
  size=10.8, after=20)
p("If you cannot define your process, you cannot measure your discipline.",
  bold=True, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# PAGE 12 — MY TRADING SYSTEM
# ---------------------------------------------------------------------

page_header(12)
title("MY TRADING SYSTEM")
for field in [
    "Market / Instrument",
    "Trading Session",
    "Timeframe Framework",
    "Market Context / Bias",
    "Entry Setup / Entry Confirmation",
    "Stop Loss Rule",
    "Take Profit / Exit Rule",
    "Risk per Trade",
    "Maximum Trades / Daily Risk Limit",
]:
    p(field, bold=True, size=9.2, after=1)
    writing_lines(1)
page_break()


# ---------------------------------------------------------------------
# PAGE 13 — MY PROCESS
# ---------------------------------------------------------------------

page_header(13)
title("MY PROCESS", "กระบวนการก่อนกด Order ของฉัน")
writing_lines(14)
page_break()


# ---------------------------------------------------------------------
# PAGE 14 — DEFINE A+ SETUP
# ---------------------------------------------------------------------

page_header(14)
title("07 — DEFINE YOUR A+ SETUP", "นิยาม Setup ที่ดีที่สุดของตัวเอง")
p("ตลาดเคลื่อนไหวตลอดเวลา แต่ไม่ได้หมายความว่าเราต้องมีส่วนร่วมกับทุก Movement",
  size=10.8, after=10)
p("A+ Setup คือสถานการณ์ที่องค์ประกอบสำคัญของระบบเรามารวมกัน "
  "และเป็น Trade ที่เราต้องการทำซ้ำเมื่อเงื่อนไขเดิมเกิดขึ้นอีก",
  size=10.8, after=13)
p("A+ Setup ไม่ได้หมายถึง Setup ที่ชนะทุกครั้ง\n"
  "แต่คือ Setup ที่ตรงตามเงื่อนไขของระบบเราอย่างชัดเจน",
  bold=True, size=11.3, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# PAGE 15 — MY A+ SETUP
# ---------------------------------------------------------------------

page_header(15)
title("MY A+ SETUP", "A+ Setup ของฉัน")
for field in [
    "Market Context",
    "Location / Zone",
    "Setup",
    "Entry Confirmation",
    "Minimum R:R",
    "สิ่งที่ทำให้ Setup นี้ Invalid",
]:
    p(field, bold=True, size=9.2, after=1)
    writing_lines(1)
page_break()


# ---------------------------------------------------------------------
# PAGE 16 — A+ EXAMPLE
# ---------------------------------------------------------------------

page_header(16)
title("MY A+ SETUP EXAMPLE", "ตัวอย่าง A+ Setup ของฉัน")
box = doc.add_table(rows=1, cols=1)
box.style = "Table Grid"
box.alignment = WD_TABLE_ALIGNMENT.CENTER
box.rows[0].height = Mm(180)
box.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
box.cell(0,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
box.cell(0,0).text = "CHART / SCREENSHOT SPACE"
box.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in box.cell(0,0).paragraphs[0].runs:
    run.font.name = FONT
    run.font.size = Pt(10)
set_cell_border(box.cell(0,0), "C4C4C4", "4")
page_break()


# ---------------------------------------------------------------------
# PAGE 17 — NO-TRADE RULES
# ---------------------------------------------------------------------

page_header(17)
title("08 — MY NO-TRADE RULES", "เมื่อไรที่ฉันจะไม่เทรด")
p("No-Trade Rules คือขอบเขตที่ช่วยป้องกันไม่ให้การตัดสินใจชั่วขณะเข้ามาแทนที่ Trading Plan",
  size=10.5, after=8)

for item in [
    "☐ ไม่มี Setup ตามระบบ",
    "☐ Entry ไม่มี Confirmation",
    "☐ Risk สูงกว่าเกณฑ์",
    "☐ R:R ไม่คุ้มกับความเสี่ยง",
    "☐ อยู่ในช่วงข่าวที่กำหนดให้หลีกเลี่ยง",
    "☐ ถึง Daily Loss Limit แล้ว",
    "☐ กำลัง FOMO",
    "☐ กำลัง Revenge Trade",
    "☐ กำลังเทรดเพราะเบื่อหรืออยากมีส่วนร่วม",
]:
    p(item, size=9.7, after=3)

p("MY PERSONAL NO-TRADE RULES — กฎ No-Trade ส่วนตัวของฉัน",
  bold=True, size=9.8, before=6, after=2)
writing_lines(3, 60)
p("NO TRADE IS ALSO A TRADING DECISION.",
  bold=True, size=10.8, align=WD_ALIGN_PARAGRAPH.CENTER, before=8)
page_break()


# ---------------------------------------------------------------------
# PAGE 18 — HOW TO USE
# ---------------------------------------------------------------------

page_header(18)
title("09 — HOW TO USE THIS JOURNAL", "วิธีใช้ Journal ตลอด 30 วัน")

flow = [
    ("PLAN", "วางแผน", "วันนี้ตลาดอยู่ใน Context แบบใด และเรากำลังรออะไร"),
    ("CHECK", "ตรวจสอบ", "Trade ที่กำลังจะเข้าอยู่ในระบบหรือไม่"),
    ("EXECUTE", "ลงมือ", "บันทึกสิ่งที่ทำจริง ไม่ใช่สิ่งที่เราคิดว่าควรทำ"),
    ("REVIEW", "ทบทวน", "ผลลัพธ์เกิดอะไรขึ้น และเราทำตามแผนหรือไม่"),
    ("REFLECT", "เรียนรู้", "วันนี้เรียนรู้อะไร และพรุ่งนี้จะปรับอะไรหนึ่งเรื่อง"),
]
for en, th, desc in flow:
    p(en, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    p(th, bold=True, size=9.8, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    p(desc, size=9.4, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

p("NO TRADE DAY COUNTS.", bold=True, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, before=6)
p("เป้าหมายไม่ใช่ 30 Days of Trading แต่คือ 30 DAYS OF PROCESS",
  bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# PAGE 19 — BEFORE DAY 1
# ---------------------------------------------------------------------

page_header(19)
title("BEFORE DAY 1", "ก่อนเริ่มต้น Day 1")

for prompt in [
    "ใน 30 วันนี้ ฉันไม่ได้ต้องการพิสูจน์ว่า...",
    "สิ่งที่ฉันต้องการเรียนรู้เกี่ยวกับตัวเองคือ...",
    "พฤติกรรมหนึ่งอย่างที่ฉันต้องการลดคือ...",
    "Process หนึ่งอย่างที่ฉันต้องการทำให้สม่ำเสมอขึ้นคือ...",
]:
    p(prompt, bold=True, size=9.8, before=5, after=1)
    writing_lines(2)

p("MY COMMITMENT — คำมั่นกับตัวเอง",
  bold=True, size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER, before=5, after=2)
p("For the next 30 days,\nI will focus on the quality of my decisions,\nnot the outcome of a single trade.",
  bold=True, size=10.3, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# DAILY JOURNAL — PAGES 20–49
# ---------------------------------------------------------------------

def daily_trade_template(day, page_no):
    page_header(page_no)

    p(f"DAY {day:02d} / 30   DAILY TRADING JOURNAL — บันทึกการเทรดประจำวัน",
      bold=True, size=13, after=4)

    meta = doc.add_table(rows=2, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(meta)
    meta_data = [
        ["DATE", "________________", "MARKET", "________________"],
        ["SESSION", "☐ Asia  ☐ London  ☐ New York  ☐ Other", "NO-TRADE DAY", "☐ Yes  ☐ No"],
    ]
    for i, row in enumerate(meta_data):
        for j, val in enumerate(row):
            meta.cell(i,j).text = val
            set_cell_margins(meta.cell(i,j), 25, 25, 25, 25)
            for run in meta.cell(i,j).paragraphs[0].runs:
                run.font.name = FONT
                run.font.size = Pt(8)
                if j % 2 == 0:
                    run.bold = True

    p("01 — PLAN | แผนก่อนเริ่มเทรด", bold=True, size=9.6, before=3, after=1)
    p("BIAS   ☐ Bullish   ☐ Bearish   ☐ Range / Neutral      KEY ZONE / LEVEL __________________",
      size=8.7, after=2)

    p("02 — PRE-TRADE CHECK | ตรวจสอบก่อนกด Order", bold=True, size=9.6, after=1)
    p("SETUP ______________________   ☐ ตรงแผน   ☐ Location เหมาะสม   ☐ Confirmation   ☐ Risk OK",
      size=8.2, after=1)
    p("☐ R:R OK   ☐ News OK   ☐ No FOMO / Revenge / Boredom", size=8.2, after=1)
    p("WHY AM I TAKING THIS TRADE?  " + "_" * 54, size=8.4, after=2)

    p("03 — TRADE LOG | บันทึกการเทรด", bold=True, size=9.6, after=1)
    p("TRADE 01  (มีกี่ไม้ กี่ Order ก็ใช้รูปแบบเดียวกัน)      ☐ BUY   ☐ SELL",
      bold=True, size=8.4, after=1)
    p("ENTRY _______________   STOP LOSS _____________   TAKE PROFIT _____________   RISK % ________",
      size=8.2, after=1)
    p("PLANNED R:R _______________     RESULT (R) _______________", size=8.2, after=2)

    p("04 — EXECUTION | ฉันทำตามแผนหรือไม่?", bold=True, size=9.6, after=1)
    p("DID I FOLLOW MY PLAN?   ☐ Yes   ☐ No      TRADE QUALITY   ☐ Good Process   ☐ Bad Process",
      size=8.2, after=1)
    p("Emotion ระหว่าง Trade มีผลต่อการตัดสินใจหรือไม่?   ☐ No   ☐ Yes   เพราะ _________________________",
      size=8.2, after=2)

    p("05 — DAILY REFLECTION | ทบทวนหลังจบวัน", bold=True, size=9.6, after=1)
    p("วันนี้ฉันทำอะไรได้ดี?  " + "_" * 56, size=8.2, after=1)
    p("วันนี้ฉันควรปรับอะไร?  " + "_" * 58, size=8.2, after=1)
    p("One thing for tomorrow  " + "_" * 55, size=8.2, after=2)

    p("PROCESS SCORE   Planning 1 2 3 4 5   Execution 1 2 3 4 5   Risk 1 2 3 4 5   Discipline 1 2 3 4 5",
      bold=True, size=7.8, after=1)
    p("TOTAL ____ /20", bold=True, size=8.2, align=WD_ALIGN_PARAGRAPH.RIGHT, after=1)
    p("DID I FOLLOW MY PROCESS TODAY?   ☐ YES   ☐ MOSTLY   ☐ NO",
      bold=True, size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


for day in range(1, 31):
    daily_trade_template(day, 19 + day)
    page_break()


# ---------------------------------------------------------------------
# PAGE 50 — END OF 30 DAYS
# ---------------------------------------------------------------------

page_header(50)
title("END OF 30 DAYS", "สิ้นสุด 30 วันแห่งการฝึก Process")
p("30 วันที่ผ่านมาอาจมีทั้งวันที่กำไร วันที่ขาดทุน วันที่ทำตามแผนได้ดี "
  "วันที่หลุดจากแผน และวันที่ไม่ได้ Trade เลย ทั้งหมดนั้นคือข้อมูล",
  size=10.8, after=12)
p("สิ่งสำคัญจึงไม่ใช่เพียงว่า “30 วันนี้ฉันทำกำไรได้เท่าไร?” "
  "แต่คือ “30 วันนี้ฉันได้เรียนรู้อะไรเกี่ยวกับตัวเองในฐานะ Trader?”",
  bold=True, size=11.2, after=15)
p("อย่ารีบตัดสินตัวเองจากผลลัพธ์\nลองกลับไปมองสิ่งที่เกิดขึ้นตลอด 30 วัน แล้วตอบจากสิ่งที่เห็นจริงใน Journal",
  size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()


# ---------------------------------------------------------------------
# PAGE 51 — LOOKING BACK
# ---------------------------------------------------------------------

page_header(51)
title("01 — LOOKING BACK", "มองย้อนกลับไป")
for prompt in [
    "สิ่งที่ฉันทำได้ดีขึ้น",
    "สิ่งที่ฉันสามารถทำตามแผนได้สม่ำเสมอขึ้น",
    "สิ่งที่ยังเป็นความท้าทายสำหรับฉัน",
]:
    p(prompt, bold=True, size=9.8, before=7, after=1)
    writing_lines(3)
page_break()


# ---------------------------------------------------------------------
# PAGE 52 — WHAT I LEARNED ABOUT MYSELF
# ---------------------------------------------------------------------

page_header(52)
title("02 — WHAT I LEARNED ABOUT MYSELF", "สิ่งที่ฉันได้เรียนรู้เกี่ยวกับตัวเอง")
for prompt in [
    "Setup หรือ Market Condition แบบไหนที่ฉันตัดสินใจได้ดีที่สุด?",
    "พฤติกรรมอะไรที่เกิดขึ้นซ้ำ ๆ?",
    "Emotion ใดมีผลต่อการตัดสินใจมากที่สุด?",
    "ฉันมักหลุดจาก Trading Plan เมื่อเกิดอะไรขึ้น?",
    "สิ่งหนึ่งที่ฉันค้นพบเกี่ยวกับตัวเองจาก Journal เล่มนี้คือ...",
]:
    p(prompt, bold=True, size=9.4, before=4, after=1)
    writing_lines(1)
page_break()


# ---------------------------------------------------------------------
# PAGE 53 — PROCESS CHECK
# ---------------------------------------------------------------------

page_header(53)
title("03 — PROCESS CHECK", "กลับมาประเมิน Process อีกครั้ง")

process_table = doc.add_table(rows=len(ASSESSMENT_ITEMS), cols=2)
process_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(process_table)
for i, item in enumerate(ASSESSMENT_ITEMS):
    process_table.cell(i,0).text = item
    process_table.cell(i,1).text = "1    2    3    4    5"
    process_table.cell(i,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for j in (0,1):
        set_cell_margins(process_table.cell(i,j), 70, 20, 70, 20)
        for run in process_table.cell(i,j).paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(9.4)
    scale_row(process_table, i, 8.8)
page_break()


# ---------------------------------------------------------------------
# PAGE 54 — DAY 1 VS DAY 30
# ---------------------------------------------------------------------

page_header(54)
title("04 — DAY 1 VS DAY 30", "ฉันเปลี่ยนไปอย่างไร?")
for prompt in [
    "สิ่งที่เปลี่ยนแปลงชัดที่สุด",
    "สิ่งที่ฉันภูมิใจกับ Process ของตัวเองมากที่สุด",
    "สิ่งที่ฉันยังต้องพัฒนาต่อ",
]:
    p(prompt, bold=True, size=9.8, before=7, after=1)
    writing_lines(3)
page_break()


# ---------------------------------------------------------------------
# PAGE 55 — STOP / START / CONTINUE
# ---------------------------------------------------------------------

page_header(55)
title("05 — STOP • START • CONTINUE", "จากวันนี้ ฉันจะ...")
for en, th in [
    ("STOP", "สิ่งที่ฉันจะหยุดทำ"),
    ("START", "สิ่งที่ฉันจะเริ่มทำ"),
    ("CONTINUE", "สิ่งที่ฉันจะทำต่อไป"),
]:
    p(en, bold=True, size=13, before=7, after=0)
    p(th, bold=True, size=9.8, after=1)
    writing_lines(3)
page_break()


# ---------------------------------------------------------------------
# PAGE 56 — MY NEXT STEP
# ---------------------------------------------------------------------

page_header(56)
title("06 — MY NEXT STEP", "ก้าวต่อไปของฉัน")

for en, th in [
    ("ONE PROCESS I WANT TO IMPROVE", "สิ่งหนึ่งที่ฉันต้องการพัฒนาต่อคือ"),
    ("WHY?", "เพราะ"),
    ("HOW?", "ฉันจะทำให้เกิดขึ้นจริงโดย"),
]:
    p(f"{en} — {th}", bold=True, size=10, before=8, after=1)
    writing_lines(3)

p("RECORD → REVIEW → RECOGNIZE → REFINE",
  bold=True, size=11.3, align=WD_ALIGN_PARAGRAPH.CENTER, before=12)
page_break()


# ---------------------------------------------------------------------
# PAGE 57 — CLOSING
# ---------------------------------------------------------------------

page_header(57)
p("30 DAYS COMPLETED", "TitleX", align=WD_ALIGN_PARAGRAPH.CENTER, before=45, after=20)
p("The goal was never perfection.\nThe goal was awareness, consistency, and improvement.",
  bold=True, size=13.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
p("เป้าหมายไม่ใช่การทำทุกอย่างให้สมบูรณ์แบบ\n"
  "แต่คือการมองเห็นตัวเอง\n"
  "สร้างความสม่ำเสมอ\n"
  "และพัฒนา Process ให้ดีขึ้นทีละน้อย",
  size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=38)
p("PROCESS BEFORE PROFIT", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
p("TRADING NOTES", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

footer()
doc.save(OUTPUT)

print(f"Created: {OUTPUT}")
