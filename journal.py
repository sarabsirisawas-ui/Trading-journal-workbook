from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
import re, os

src="/mnt/data/Trading_Notes_Daily_Journal_30_Days_A4_Book_Rebalanced.docx"
out="/mnt/data/Trading_Notes_Daily_Journal_30_Days_A4_Book_Layout_v2.docx"
doc=Document(src)

# Make the foundation/closing feel like a real printed book:
# larger reading size + much more vertical rhythm, while daily pages remain practical.
in_daily=False
for par in doc.paragraphs:
    txt=par.text.strip()
    if re.match(r"^DAY 01 / 30$", txt): in_daily=True
    if txt=="END OF 30 DAYS": in_daily=False

    if not in_daily:
        # Major English headings
        if re.match(r"^\d{2}\s*[—-]", txt) or txt in {
            "CONTENTS","DISCLAIMER","THIS WORKBOOK BELONGS TO","SELF-ASSESSMENT REFLECTION",
            "MY TRADING SYSTEM","MY A+ SETUP","BEFORE DAY 1","END OF 30 DAYS",
            "30 DAYS COMPLETED"
        }:
            for r in par.runs:
                r.font.size=Pt(22)
                r.bold=True
            par.paragraph_format.space_before=Pt(12)
            par.paragraph_format.space_after=Pt(6)

        # Thai subheadings / secondary headings
        elif txt and len(txt)<75 and any(k in txt for k in [
            "ก่อนเริ่มต้น","ทำไม Trader","กระบวนการมาก่อน","Trade ที่ดี","ก่อนวิเคราะห์",
            "นิยามระบบ","นิยาม Setup","เมื่อไรที่ฉัน","วิธีใช้ Journal","Workbook เล่มนี้",
            "ข้อจำกัด","สารบัญ","ทบทวนตัวเอง","ระบบการเทรดของฉัน","A+ Setup ของฉัน",
            "ก่อนเริ่มต้น Day 1","สิ้นสุด 30 วัน","มองย้อนกลับไป","สิ่งที่ฉันได้เรียนรู้",
            "กลับมาประเมิน","ฉันเปลี่ยนไป","จากวันนี้ ฉันจะ","ก้าวต่อไปของฉัน"
        ]):
            for r in par.runs:
                r.font.size=Pt(14.5)
                r.bold=True
            par.paragraph_format.space_after=Pt(13)

        # Main prose
        elif len(txt)>55 and "_" not in txt:
            for r in par.runs:
                r.font.size=Pt(12.6)
            par.paragraph_format.line_spacing=1.42
            par.paragraph_format.space_after=Pt(12)

        # Short emphasized lines
        elif txt and "_" not in txt and len(txt)<=55:
            for r in par.runs:
                if r.font.size is None or r.font.size.pt < 10:
                    r.font.size=Pt(11.3)
            par.paragraph_format.space_after=Pt(7)

        # Writing lines: spread them vertically
        if txt.count("_")>40:
            for r in par.runs:
                r.font.size=Pt(10.2)
            par.paragraph_format.line_spacing=1.65
            par.paragraph_format.space_after=Pt(11)

# Specific sparse pages: create more generous vertical pacing without adding content
for par in doc.paragraphs:
    txt=par.text.strip()
    if txt in ("If you cannot define your process, you cannot measure your discipline.",
               "NO TRADE IS ALSO A TRADING DECISION.",
               "Don't judge the process by one outcome."):
        par.paragraph_format.space_before=Pt(26)
        par.paragraph_format.space_after=Pt(16)
        for r in par.runs:
            r.font.size=Pt(13)
            r.bold=True

# Foundation tables: bigger, more editorial
for table in doc.tables:
    rows,cols=len(table.rows),len(table.columns)
    # skip daily tables by recognizing their structures but enlarge only non-daily matrix / chart
    if rows==3 and cols==3:
        for row in table.rows:
            row.height=Mm(15)
            row.height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.font.size=Pt(11)
    if rows==1 and cols==1 and "CHART / SCREENSHOT SPACE" in table.cell(0,0).text:
        table.rows[0].height=Mm(88)
        table.rows[0].height_rule=WD_ROW_HEIGHT_RULE.EXACTLY

doc.save(out)
print(out)
